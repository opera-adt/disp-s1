#!/usr/bin/env python
from __future__ import annotations

import argparse
from itertools import groupby

import h5py
from docx import Document
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

Filename = str  # or a Path type, etc.

GROUP_NAME_MAP = {
    "root": "Imagery layers",
    "corrections": "Corrections",
    "identification": "Identification",
    "metadata": "Metadata",
}

MID_GRAY_COLOR = "a3a3a3"
LIGHT_GRAY_COLOR = "EFEFEF"


def _pretty_group_name(hdf5_path: str) -> str:
    """Derive top-level group name from the dataset path, map to user-friendly name."""
    parts = hdf5_path.strip("/").split("/")
    group = parts[-2] if len(parts) > 1 else "root"
    return GROUP_NAME_MAP.get(group, group)


def _fmt_dtype(dtype) -> str:
    """Return a short descriptive string for HDF5/numpy dtypes."""
    kind = getattr(dtype, "kind", None)
    if kind in ("S", "O", "U"):
        return "string"
    elif kind == "f":
        return f"float{dtype.itemsize * 8}"
    elif kind == "i":
        return f"int{dtype.itemsize * 8}"
    elif kind == "u":
        return f"uint{dtype.itemsize * 8}"
    return str(dtype)


def _fmt_shape(shape) -> str:
    """Format shape as 'scalar', '1D', '2D', or 'ND'."""
    if not shape or shape == ():
        return "scalar"
    elif len(shape) == 1:
        return "1D"
    elif len(shape) == 2:
        return "2D"
    else:
        return f"{len(shape)}D"  # or just "ND"


def _to_str(item) -> str:
    """Ensure we always return a Python str from either str or bytes."""
    if isinstance(item, bytes):
        return item.decode("utf-8")
    return str(item)


def _get_table_data(hdf5_path: Filename):
    """Collect dataset info (Name, Group, Type, Shape, Units, Description)."""
    table_data = []

    def append_dset(name, item):
        if not isinstance(item, h5py.Dataset):
            return
        dtype_str = _fmt_dtype(item.dtype)
        shape_str = _fmt_shape(item.shape)
        # Prefer 'description' if present, else 'long_name', else 'No description'
        desc = item.attrs.get("description", b"") or item.attrs.get("long_name", b"")
        desc = desc if desc else b"No description"
        units = item.attrs.get("units", b"")

        table_data.append(
            {
                "Name": name,
                "Group": _pretty_group_name(name),
                "Type": dtype_str,
                "Shape": shape_str,
                "Units": _to_str(units),
                "Description": _to_str(desc),
            }
        )

    with h5py.File(hdf5_path, "r") as hf:
        hf.visititems(append_dset)

    # Group by "Group" key (the friendly group name)
    table_data.sort(key=lambda x: x["Group"])
    groups = {}
    for g, rows in groupby(table_data, key=lambda x: x["Group"]):
        groups[g] = list(rows)
    return groups


def generate_docx_table(hdf5_path: Filename, output_path: Filename):
    """Create a .docx file with a table for `hdf5_path` metadata.

    Parameters
    ----------
    hdf5_path : Filename
        Input HDF5 file
    output_path : Filename
        Output .docx file to save results.

    """
    document = Document()
    # Default font
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"

    groups = _get_table_data(hdf5_path)

    # We'll iterate in a fixed order:
    for group_name in ["Imagery layers", "Corrections", "Identification", "Metadata"]:
        rows = groups[group_name]
        document.add_heading(group_name, level=2)

        # Create a table with 4 columns for: Name, Type, Shape, Units
        # We'll place the Description in a second row each time.
        table = document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        table.preferred_width = Inches(6.16)

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Type"
        hdr_cells[2].text = "Shape"
        hdr_cells[3].text = "Units"

        # Gray shading on header
        for cell in hdr_cells:
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{MID_GRAY_COLOR}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # Populate rows
        for row_dict in rows:
            # First row: name/type/shape/units
            row = table.add_row()
            row_cells = row.cells
            # Set row height to 0.35" (At Least)
            # table_row = row[0]._tc.getparent()
            # table_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            # table_row.height = Inches(0.35)

            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.35)

            dataset_shortname = row_dict["Name"].split("/")[-1]
            row_cells[0].text = _to_str(dataset_shortname)
            row_cells[0].width_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row_cells[1].text = _to_str(row_dict["Type"])
            row_cells[2].text = _to_str(row_dict["Shape"])
            units = _to_str(row_dict["Units"])
            if units.startswith("seconds since"):
                units = "seconds"  # Caveat about reference epoch to be made in document
            row_cells[3].text = units

            # Second row: single cell spanning all 4 columns, with Description
            desc_row = table.add_row()
            desc_row_cells = desc_row.cells
            desc_row_cells[0].merge(desc_row_cells[1]).merge(desc_row_cells[2]).merge(
                desc_row_cells[3]
            )
            desc_row_cells[0].text = row_dict["Description"]
            desc_row_cells[0].shading = WD_BUILTIN_STYLE.TABLE_LIGHT_SHADING
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY_COLOR}"/>'
            )
            desc_row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
            # desc_table_row = desc_row[0]._tc.getparent()
            desc_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            desc_row.height = Inches(0.35)

        # widths = np.array([1.5, 0.35, 0.2, 0.8])
        # widths / widths.sum() * 6.1
        # array([2.77272727, 0.92424242, 0.92424242, 1.47878788])

        # Optional column width
        # Name, Type, Shape, Units
        # col_widths = [Inches(1.5), Inches(0.5), Inches(0.5), Inches(0.8)]
        col_widths = [Inches(3.2), Inches(0.7), Inches(0.4), Inches(1.7)]
        for i, width in enumerate(col_widths):
            for r in table.rows:
                r.cells[i].width = width

    document.save(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("hdf5_path", help="Path to the HDF5 file")
    parser.add_argument("output_path", help="Path to the output Word docx file")
    args = parser.parse_args()
    generate_docx_table(args.hdf5_path, args.output_path)
